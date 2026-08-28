import io
import logging

import docx
import pymupdf

from app.core.text_utils import (
    EMAIL_PATTERN,
    PHONE_PATTERN,
    URL_PATTERN,
    count_words,
    normalize_whitespace,
)
from app.schemas.parsing import ContactInfo, LayoutFindings, ParsedResume

logger = logging.getLogger(__name__)


class UnsupportedFileTypeError(Exception):
    pass


class EmptyResumeError(Exception):
    pass


def _detect_file_kind(file_name: str, mime_type: str) -> str | None:
    extension = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
    if extension == "pdf" or "pdf" in mime_type:
        return "pdf"
    if extension == "docx" or "wordprocessingml" in mime_type:
        return "docx"
    if extension == "txt" or mime_type.startswith("text/plain"):
        return "txt"
    return None


def _detect_multi_column(page: "pymupdf.Page") -> bool:
    """
    Best-effort heuristic: true side-by-side columns show up as two sets of
    blocks confined to the left/right halves of the page whose vertical
    (y) ranges overlap - i.e. text running in parallel, not just a header
    photo or an isolated pull-quote.
    """
    blocks = page.get_text("blocks")
    page_width = page.rect.width
    mid = page_width / 2

    left_ranges: list[tuple[float, float]] = []
    right_ranges: list[tuple[float, float]] = []
    for block in blocks:
        x0, y0, x1, y1, text = block[0], block[1], block[2], block[3], block[4]
        if not text.strip() or (x1 - x0) < 20:
            continue
        if x1 <= mid + 10:
            left_ranges.append((y0, y1))
        elif x0 >= mid - 10:
            right_ranges.append((y0, y1))

    if len(left_ranges) < 2 or len(right_ranges) < 2:
        return False

    overlap_count = sum(
        1
        for ly0, ly1 in left_ranges
        if any(ly0 < ry1 and ry0 < ly1 for ry0, ry1 in right_ranges)
    )
    return overlap_count >= 2


def _parse_pdf(buffer: bytes) -> tuple[str, LayoutFindings]:
    doc = pymupdf.open(stream=buffer, filetype="pdf")
    try:
        page_texts: list[str] = []
        table_count = 0
        image_count = 0
        multi_column = False

        for page in doc:
            page_texts.append(page.get_text("text", sort=True))
            try:
                tables = page.find_tables()
                table_count += len(tables.tables)
            except Exception:
                logger.debug("Table detection failed for a page", exc_info=True)
            image_count += len(page.get_images(full=True))
            if not multi_column:
                multi_column = _detect_multi_column(page)

        text = "\n\n".join(page_texts)
        layout = LayoutFindings(
            page_count=doc.page_count,
            has_tables=table_count > 0,
            table_count=table_count,
            has_images=image_count > 0,
            image_count=image_count,
            likely_multi_column=multi_column,
        )
        return text, layout
    finally:
        doc.close()


def _parse_docx(buffer: bytes) -> tuple[str, LayoutFindings]:
    document = docx.Document(io.BytesIO(buffer))
    paragraphs = [p.text for p in document.paragraphs]
    text = "\n".join(paragraphs)

    layout = LayoutFindings(
        page_count=None,
        has_tables=len(document.tables) > 0,
        table_count=len(document.tables),
        has_images=len(document.inline_shapes) > 0,
        image_count=len(document.inline_shapes),
        likely_multi_column=False,  # not detectable from python-docx's model
    )
    return text, layout


def _extract_contact_info(text: str) -> ContactInfo:
    emails = EMAIL_PATTERN.findall(text)
    urls = URL_PATTERN.findall(text)
    return ContactInfo(
        has_email=len(emails) > 0,
        has_phone=bool(PHONE_PATTERN.search(text)),
        has_url=len(urls) > 0,
        emails=emails[:5],
        urls=urls[:5],
    )


def parse_resume(buffer: bytes, file_name: str, mime_type: str) -> ParsedResume:
    kind = _detect_file_kind(file_name, mime_type)
    if kind is None:
        raise UnsupportedFileTypeError(
            f'Unsupported file type for "{file_name}". Please upload a PDF, DOCX, or TXT file.'
        )

    if kind == "pdf":
        raw_text, layout = _parse_pdf(buffer)
    elif kind == "docx":
        raw_text, layout = _parse_docx(buffer)
    else:
        raw_text = buffer.decode("utf-8", errors="replace")
        layout = LayoutFindings()

    text = normalize_whitespace(raw_text)
    if len(text) < 30:
        raise EmptyResumeError(
            "We couldn't find any readable text in this file. If it's a scanned "
            "image, try exporting a text-based PDF or DOCX instead."
        )

    return ParsedResume(
        text=text,
        word_count=count_words(text),
        file_kind=kind,
        contact=_extract_contact_info(text),
        layout=layout,
    )
